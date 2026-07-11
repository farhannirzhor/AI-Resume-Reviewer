import io

import pdfplumber
from docx import Document
from fastapi import HTTPException, UploadFile

ALLOWED_EXTENSIONS = {".pdf", ".docx"}


def _get_extension(filename: str) -> str:
    if "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    text_chunks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks).strip()


def _extract_text_from_docx(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    text_chunks = [para.text for para in document.paragraphs if para.text.strip()]

    # also pull text out of tables, since resumes sometimes use them for layout
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_chunks.append(cell.text.strip())

    return "\n".join(text_chunks).strip()


async def extract_resume_text(file: UploadFile, max_size_mb: int) -> str:
    """Validate the uploaded resume file and return its extracted plain text."""
    extension = _get_extension(file.filename or "")

    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Please upload a PDF or DOCX resume.",
        )

    file_bytes = await file.read()

    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > max_size_mb:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Max allowed size is {max_size_mb} MB.",
        )

    try:
        if extension == ".pdf":
            text = _extract_text_from_pdf(file_bytes)
        else:
            text = _extract_text_from_docx(file_bytes)
    except Exception:
        raise HTTPException(
            status_code=400,
            detail="Could not read the uploaded file. It may be corrupted or scanned as an image.",
        )

    if not text or len(text) < 50:
        raise HTTPException(
            status_code=400,
            detail="Could not extract enough text from the resume. "
                   "If this is a scanned/image-based PDF, please upload a text-based file instead.",
        )

    return text
